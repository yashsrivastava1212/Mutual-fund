"""Convert InterviewPreparation.md to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DEFAULT_INPUT = PROJECT_ROOT / "content" / "InterviewPreparation.md"
DEFAULT_OUTPUT = PROJECT_ROOT / "content" / "InterviewPreparation.docx"


def _add_formatted_run(paragraph, text: str, *, bold: bool = False, italic: bool = False, code: bool = False) -> None:
    """Add text with optional inline **bold** and `code` parsing."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells if c)


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        if _is_table_row(stripped):
            table_lines: list[str] = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [_parse_table_row(row) for row in table_lines]
            rows = [r for r in rows if not _is_separator_row(r)]
            if rows:
                col_count = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(col_count):
                        text = row[c_idx] if c_idx < len(row) else ""
                        table.rows[r_idx].cells[c_idx].text = text
                doc.add_paragraph()
            continue

        if stripped.startswith("| ") and "---" in stripped:
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        if stripped.startswith("**Q") and ":" in stripped:
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped, bold=True)
            i += 1
            continue

        if stripped.startswith("**A:") or stripped.startswith("**Follow-up"):
            p = doc.add_paragraph()
            _add_formatted_run(p, stripped)
            i += 1
            continue

        p = doc.add_paragraph()
        _add_formatted_run(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_INPUT
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT
    convert(input_path, output_path)
